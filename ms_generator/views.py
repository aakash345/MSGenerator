import os
import shutil
from django.shortcuts import render, HttpResponse, redirect
import pandas as pd
from docx2pdf import convert
from django.conf.urls.static import static
from docx import Document
import win32com.client
import pythoncom
from task.settings import MEDIA_ROOT
# Create your views here.

def home(request):
    if request.method=="POST":
        wpath = MEDIA_ROOT+ "/word/"
        ppath = MEDIA_ROOT+ "/output/"
        for file_name in os.listdir(wpath):
            # construct full file path
            file = wpath + file_name
            if os.path.isfile(file):
                print('Deleting file:', file)
                os.remove(file)
        for file_name in os.listdir(ppath):
            # construct full file path
            file = ppath + file_name
            if os.path.isfile(file):
                print('Deleting file:', file)
                os.remove(file)
        os.remove(MEDIA_ROOT+ "/result.zip")
        return render(request, "home.html")
    return render(request, 'home.html')

def output(request):
    if request.method=="POST":
        user_in = request.FILES['marksdat']
        user_dat = request.FILES['basicdat']
        user_in = pd.read_excel(user_in)
        user_dat = pd.read_excel(user_dat)
        mydic = {"ENG":"English","MATH":"Engineering Mathematics","BME":"Basics of Mechanical Engineering","BOE":"Basics of Electronics","PHY":"Engineering Physics","EEE":"Electrical Engineering","DSA":"Data Structure and Algorithm","DT":"Data Transformation","EVS":"Environmental Studies"}
        f = Document(MEDIA_ROOT + "/mtemp/template.docx")
        table1 = f.tables[0]
        table2 = f.tables[1]
        table3 = f.tables[2]
        pythoncom.CoInitialize()
        for ind in user_in.index:
            total_marks = 0
            marks_obtained = 0
            grade = ""
            user_roll = user_in['roll'][ind]
            user = user_dat[user_dat['roll'] == user_roll]
            name = user['name'].to_string().split()[1]
            table1.cell(0, 1).text = user['name']
            table1.cell(0, 3).text = user['roll']
            table1.cell(1, 1).text = user['course']
            table1.cell(1, 3).text = user['branch']
            table1.cell(2, 1).text = user['parentName']
            user = user_in[user_in['roll'] == user_roll]
            for i in range(1,7):
                table2.cell(i, 0).text = f'{i}'
                sub = user[f'subject{i}']
                marks = user[f'markObSub{i}']
                marks_obtained += int(marks)
                maxmarks = user[f'maxMarkSub{i}']
                total_marks += int(maxmarks)
                table2.cell(i, 1).text = sub
                table2.cell(i, 2).text = mydic[sub.to_string().split()[1]]
                table2.cell(i, 3).text = marks.to_string().split()[1]
                table2.cell(i, 4).text = maxmarks.to_string().split()[1]
            print(total_marks)
            print(marks_obtained)
            percent = (marks_obtained/total_marks)*100
            if(percent>=90):
                grade = "A"
            elif(percent>=80):
                grade = "B"
            elif(percent>=70):
                grade = "C"
            elif(percent>=60):
                grade = "D"
            elif(percent>=40):
                grade = "E"
            else:
                grade = "F"
            print(grade)
            table3.cell(0, 1).text = str(marks_obtained)
            table3.cell(1, 1).text = str(total_marks)
            table3.cell(2, 1).text = grade
            f.save(MEDIA_ROOT+ f"/word/{name} {user_roll}.docx")
        
        convert(MEDIA_ROOT+"/word/", MEDIA_ROOT+"/output/")
        shutil.make_archive(MEDIA_ROOT+"/result", "zip", MEDIA_ROOT+ "\output\ ")
        return render(request, 'download.html')
